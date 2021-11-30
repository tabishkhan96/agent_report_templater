from abc import ABC, abstractmethod
from dataclasses import dataclass
from typing import List, Union, Any, Iterator, BinaryIO, Final

import docx
from docx.document import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.table import Table as DocxTable
from docx.shared import Cm

from src.repository.models import Table, Cell

# TODO возвращать объекты из pydocx в адаптерах


@dataclass
class Style:
    alignment: str
    italic: bool
    bold: bool
    font: str


DEFAULT_STYLE: Final[Style] = Style(alignment='center', italic=False, bold=True, font="Times New Roman")


class DocumentDAOInterface(ABC):
    """Интерфейс класса доступа к документам"""
    def __init__(self, path: str):
        self._document: Union[DocxDocument, Any] = self.load(path)

    @abstractmethod
    def load(self, path: str):
        """Load document form disc"""

    @abstractmethod
    def load_from_storage(self, guid: str):
        """Load document form storage"""

    @abstractmethod
    def save(self, doc_name: str):
        """Save document to storage"""

    @abstractmethod
    def get_tables(self) -> Iterator[Table]:
        """Get list of tables of Doc"""

    @abstractmethod
    def append_table(self, table: Table) -> Table:
        """Add Table to the end of Doc"""

    @abstractmethod
    def get_paragraphs(self) -> List[str]:
        """Get list of paragraphs"""

    @abstractmethod
    def append_paragraph(self, text: str, style: Style = DEFAULT_STYLE):
        """Add text paragraph to the end of Doc"""

    @abstractmethod
    def append_picture(self, picture: BinaryIO):
        """Add a picture to the end of Doc"""

    @abstractmethod
    def add_page_break(self):
        """Add page break of Doc"""

    @abstractmethod
    def add_section(self, horizontal: bool = False):
        """"""

    @classmethod
    @abstractmethod
    def set_cell_style(cls, cell: Cell, style: Style = DEFAULT_STYLE):
        """Set table cell style shortcut"""

    @classmethod
    @abstractmethod
    def insert_picture_into_cell(cls, cell: Cell, pic: BinaryIO, height: int = 8, width: int = 8):
        """Insert picture into cell shortcut"""


class DocxDocumentDAO(DocumentDAOInterface):
    """Класс доступа к документам типа .docx """
    def load(self, path: str):
        """Load document form disc"""
        return docx.Document(path)

    def load_from_storage(self, guid: str):
        """Load document form storage"""
        raise NotImplementedError

    def save(self, doc_name: str):
        """Save document to storage"""
        self._document.save(doc_name)

    def get_tables(self) -> Iterator[DocxTable]:
        """Get list of tables of Doc"""
        for table in self._document.tables:
            yield table

    def append_table(self, table: DocxTable) -> Table:
        """Add Table to the end of Doc"""
        tbl = table._tbl
        paragraph = self._document.add_paragraph()
        paragraph._p.addnext(tbl)
        return self._document.tables[-1]

    def get_paragraphs(self) -> List[str]:
        """"Get list of paragraphs"""
        return [paragraph.text for paragraph in self._document.paragraphs]

    def append_paragraph(self, text: str, style: Style = DEFAULT_STYLE):
        paragraph = self._document.add_paragraph(text=text)
        paragraph.alignment = getattr(WD_PARAGRAPH_ALIGNMENT, style.alignment.upper())
        paragraph.runs[0].bold = style.bold
        paragraph.runs[0].italic = style.italic
        paragraph.runs[0].font.name = style.font

    def append_picture(self, picture: BinaryIO, height: int = 8, width: int = 8):
        self._document.add_paragraph().add_run().add_picture(picture, width=Cm(width), height=Cm(height))

    def add_page_break(self):
        self._document.add_page_break()

    def add_section(self, horizontal: bool = False):
        section = self._document.add_section()
        if horizontal:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

    @classmethod
    def set_cell_style(cls, cell: Cell, style: Style = DEFAULT_STYLE):
        cell.paragraphs[0].paragraph_format.alignment = getattr(WD_TABLE_ALIGNMENT, style.alignment.upper())
        cell.paragraphs[0].runs[0].bold = style.bold
        cell.paragraphs[0].runs[0].italic = style.italic
        cell.paragraphs[0].runs[0].font.name = style.font

    @classmethod
    def insert_picture_into_cell(cls, cell: Cell, pic: BinaryIO, height: int = 8, width: int = 8):
        cell.paragraphs[0].add_run().add_picture(pic, width=Cm(width), height=Cm(height))

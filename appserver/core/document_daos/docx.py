from typing import List, BinaryIO, Iterator

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.shared import Cm

from .abstract import AbstractDocumentDAO, BaseAdapter, Table, Row, Column, Cell, Style, DEFAULT_STYLE


class DocxCellAdapter(Cell, BaseAdapter):
    """PyDocx table cell adapter class"""
    @property
    def text(self):
        return self._source.text

    @text.setter
    def text(self, value: str):
        self._source.text = value

    @property
    def paragraphs(self) -> Iterator:
        for paragraph in self._source.paragraphs:
            yield paragraph


class DocxRowAdapter(Row, BaseAdapter):
    """PyDocx table row adapter class"""
    @property
    def height(self) -> float:
        return self._source.height.cm

    @property
    def cells(self) -> List[Cell]:
        return [DocxCellAdapter(cell) for cell in self._source.cells]


class DocxColumnAdapter(Column, BaseAdapter):
    """PyDocx table column adapter class"""
    @property
    def width(self) -> float:
        return self._source.width.cm

    @property
    def cells(self) -> List[Cell]:
        return [DocxCellAdapter(cell) for cell in self._source.cells]


class DocxTableAdapter(Table, BaseAdapter):
    """PyDocx table adapter class"""
    @property
    def columns(self) -> list[Column]:
        return [DocxColumnAdapter(column) for column in self._source.columns]

    @property
    def rows(self) -> list[Row]:
        return [DocxRowAdapter(row) for row in self._source.rows]

    def add_row(self) -> Row:
        return DocxRowAdapter(self._source.add_row())

    def column_cells(self, column_number: int) -> Iterator[Cell]:
        for cell in self._source.column_cells(column_number):
            yield DocxCellAdapter(cell)

    def row_cells(self, row_number: int) -> Iterator[Cell]:
        for cell in self._source.row_cells(row_number):
            yield DocxCellAdapter(cell)


class DocxDocumentDAO(AbstractDocumentDAO):
    """.docx documents access class"""
    def load(self, path: str) -> docx.Document:
        """Load document form disc"""
        return docx.Document(path)

    def save(self, doc_name: str):
        """Save document to storage"""
        self._document.save(doc_name)

    def get_tables(self) -> Iterator[Table]:
        """Get list of tables of Doc"""
        for table in self._document.tables:
            yield DocxTableAdapter(table)

    def append_table(self, table: BaseAdapter) -> Table:
        """Add Table to the end of the Doc"""
        tbl = table._source._tbl
        paragraph = self._document.add_paragraph()
        paragraph._p.addnext(tbl)
        return DocxTableAdapter(self._document.tables[-1])

    def get_paragraphs(self) -> List[str]:
        """"Get list of paragraphs"""
        return [paragraph.text for paragraph in self._document.paragraphs]

    def append_paragraph(self, text: str, style: Style = DEFAULT_STYLE):
        paragraph = self._document.add_paragraph(text=text)
        paragraph.alignment = getattr(WD_PARAGRAPH_ALIGNMENT, style.alignment.upper())
        paragraph.runs[0].bold = style.bold
        paragraph.runs[0].italic = style.italic
        paragraph.runs[0].font.name = style.font

    def append_picture(self, picture: BinaryIO, height: float, width: float, alignment: str = 'center'):
        paragraph = self._document.add_paragraph()
        paragraph.alignment = getattr(WD_PARAGRAPH_ALIGNMENT, alignment.upper())
        paragraph.add_run().add_picture(picture, width=Cm(width), height=Cm(height))

    def add_page_break(self):
        self._document.add_page_break()

    def add_section(self, horizontal: bool = False):
        section = self._document.add_section()
        if horizontal:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

    def get_page_size(self) -> tuple[float, float]:
        section = self._document.sections[-1]
        return section.page_height.cm, section.page_width.cm

    @classmethod
    def set_cell_style(cls, cell: Cell, style: Style = DEFAULT_STYLE):
        paragraph = next(cell.paragraphs)
        paragraph.paragraph_format.alignment = getattr(WD_TABLE_ALIGNMENT, style.alignment.upper())
        paragraph.runs[0].bold = style.bold
        paragraph.runs[0].italic = style.italic
        paragraph.runs[0].font.name = style.font

    @classmethod
    def insert_picture_into_cell(cls, cell: Cell, pic: BinaryIO, height: float, width: float):
        next(cell.paragraphs).add_run().add_picture(pic, width=Cm(width), height=Cm(height))
